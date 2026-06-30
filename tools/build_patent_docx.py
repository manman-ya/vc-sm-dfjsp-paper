from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(
    r"D:/Users/lcp/Documents/xwechat_files/wxid_paqodkiv41g322_26f8/msg/file/2026-05/一种基于半监督随机配置网络的故障诊断方法20251124.docx"
)
OUT_DIR = ROOT / "patent"
OUT_DOCX = OUT_DIR / "一种面向多服务价值链协同的共享制造分布式柔性作业车间多目标调度方法及系统_专利稿.docx"


TITLE = "一种面向多服务价值链协同的共享制造分布式柔性作业车间多目标调度方法及系统"


def clear_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def set_east_asia_font(run, name: str = "宋体", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_para(doc: Document, text: str = "", *, bold: bool = False, size: int = 12, align=None, first_line: bool = True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if first_line and text and not bold and not text.startswith(("S", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.", "图")):
        p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    set_east_asia_font(r, size=size, bold=bold)
    return p


def add_heading(doc: Document, text: str):
    return add_para(doc, text, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def add_section_title(doc: Document, text: str):
    return add_para(doc, text, bold=True, size=13, first_line=False)


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)


def add_formula(doc: Document, text: str):
    p = add_para(doc, text, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    for run in p.runs:
        run.font.name = "Cambria Math"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_claims(doc: Document):
    add_heading(doc, "权利要求书")
    claims = [
        (
            "1. 一种面向多服务价值链协同的共享制造分布式柔性作业车间多目标调度方法，其特征在于，"
            "通过第三方共享制造平台对多条服务价值链的订单、制造服务资源单元及其内部机器进行统一建模和求解，"
            "在满足订单服务类型匹配、工艺顺序约束和机器加工约束的条件下，同时优化总成本、最大完工时间和资源负载；"
            "所述方法包括以下步骤："
        ),
        "S1、获取待调度订单集合、服务价值链集合、制造服务类型集合、制造服务资源单元集合以及每个制造服务资源单元内部的机器集合，并获取各订单的价值链归属、服务类型、工序序列、可选机器、加工时间、加工成本、运输时间、运输成本和跨链协同成本；",
        "S2、根据订单服务类型和跨链调度模式构造每个订单的候选制造服务资源单元集合，将候选制造服务资源单元划分为链内候选集合和跨链候选集合；",
        "S3、构造调度个体的四层编码，所述四层编码包括订单到制造服务资源单元的分配层UA、按服务类型组织的工序排序层OS、由UA和OS推导得到的制造服务资源单元内工序序列层OP，以及每道工序对应机器选择的机器层MS；",
        "S4、对四层编码执行可行性修复，使每个订单仅分配至一个服务类型匹配的制造服务资源单元，并使工序排序层满足各订单工序出现次数要求、机器层满足所选制造服务资源单元内可加工机器要求；",
        "S5、根据修复后的四层编码进行解码，依次计算每道工序的开始时间、结束时间、加工机器、加工成本、运输成本、跨链协同成本、各制造服务资源单元负载和订单完工时间；",
        "S6、基于步骤S5得到的评价结果形成多目标函数值，所述多目标函数至少包括总成本和最大完工时间，并能够进一步包括最大制造服务资源单元负载；",
        "S7、利用非支配排序、拥挤距离和非支配解记忆池筛选优质调度个体，并根据优质调度个体更新UA、OS和MS对应的概率模型；",
        "S8、基于更新后的概率模型采样生成新的调度个体，并结合链内资源替换、跨链资源替换、跨链回流、工序插入、机器替换和瓶颈资源释放邻域进行局部搜索；",
        "S9、迭代执行步骤S4至S8直至满足停止条件，输出非支配调度方案集合或从非支配调度方案集合中选择折中调度方案。",
        (
            "2. 根据权利要求1所述的方法，其特征在于，所述候选制造服务资源单元集合的构造方式为："
            "对于订单j，若制造服务资源单元u的可服务类型集合包含订单j的服务类型，则将制造服务资源单元u加入订单j的候选集合；"
            "若制造服务资源单元u与订单j属于同一服务价值链，则将其加入链内候选集合；若二者属于不同服务价值链且跨链调度模式允许，则将其加入跨链候选集合。"
        ),
        (
            "3. 根据权利要求1所述的方法，其特征在于，步骤S4中的可行性修复包括："
            "当UA层将订单分配至不满足服务类型匹配或跨链模式约束的制造服务资源单元时，从该订单的可行候选集合中重新选择制造服务资源单元；"
            "当OS层中订单工序令牌数量与订单工序数量不一致时，增补或删除对应令牌；"
            "当MS层所选机器不能加工对应工序时，从所选制造服务资源单元内部的可行机器中重新选择机器。"
        ),
        (
            "4. 根据权利要求1所述的方法，其特征在于，步骤S5中的解码方式为："
            "按照OP层中各制造服务资源单元内的工序序列，结合MS层机器选择，维护订单就绪时间和机器就绪时间；"
            "每道工序的开始时间取订单就绪时间和对应机器就绪时间的较大值，结束时间为开始时间加加工时间；"
            "解码过程中同步更新订单就绪时间、机器就绪时间、制造服务资源单元负载和加工成本。"
        ),
        (
            "5. 根据权利要求1所述的方法，其特征在于，步骤S6中的总成本由加工成本、运输成本和跨链固定协同成本构成，"
            "其表达式为F1=PC+TC+CFC，其中PC为所有工序加工时间与单位加工成本乘积之和，TC为订单选择对应制造服务资源单元产生的运输成本之和，"
            "CFC为订单跨服务价值链选择制造服务资源单元产生的固定协同成本之和；最大完工时间为F2=max(Cj+ttj,u)，其中Cj为订单j最后一道工序的结束时间，ttj,u为订单j选择制造服务资源单元u产生的运输时间。"
        ),
        (
            "6. 根据权利要求5所述的方法，其特征在于，所述多目标函数还包括最大制造服务资源单元负载F3=max(loadu)，"
            "其中loadu为制造服务资源单元u承接的所有工序加工时间之和。"
        ),
        (
            "7. 根据权利要求1所述的方法，其特征在于，步骤S7中的概率模型包括："
            "用于学习订单到制造服务资源单元分配偏好的PMA概率矩阵，用于学习工序排序位置偏好的PMS概率矩阵，"
            "以及用于学习工序机器选择偏好的PMM概率矩阵；每轮迭代中，以非支配解和精英解作为学习样本，按预设学习率对PMA、PMS和PMM进行平滑更新。"
        ),
        (
            "8. 根据权利要求1所述的方法，其特征在于，步骤S8中的局部搜索包括以下至少一种邻域："
            "将订单在同一服务价值链内同服务类型制造服务资源单元之间迁移的链内替换邻域；"
            "将订单迁移至其他服务价值链中同服务类型制造服务资源单元的跨链替换邻域；"
            "将跨链订单迁回其所属服务价值链内制造服务资源单元的跨链回流邻域；"
            "从当前负载最高的制造服务资源单元中选择订单迁移至低负载候选制造服务资源单元的瓶颈资源释放邻域；"
            "以及在工序排序层中移动工序令牌的插入邻域和在机器层中替换加工机器的机器替换邻域。"
        ),
        (
            "9. 一种应用如权利要求1至8任一项所述方法的共享制造多目标调度系统，其特征在于，包括："
            "数据获取模块，用于获取订单、服务价值链、服务类型、制造服务资源单元、机器、加工时间、加工成本、运输成本和跨链协同成本；"
            "候选资源构造模块，用于构造订单的链内候选制造服务资源单元集合和跨链候选制造服务资源单元集合；"
            "编码初始化模块，用于生成UA、OS、OP和MS四层编码；"
            "可行性修复模块，用于修复服务类型不匹配、跨链模式不满足、工序令牌不一致和机器不可加工的编码；"
            "解码评价模块，用于计算调度方案的开始时间、结束时间、总成本、最大完工时间和制造服务资源单元负载；"
            "概率学习模块，用于根据精英解和非支配解更新概率模型；"
            "局部搜索模块，用于执行链内替换、跨链替换、跨链回流、瓶颈资源释放、工序插入和机器替换；"
            "输出模块，用于输出非支配调度方案集合或折中调度方案。"
        ),
        (
            "10. 一种计算机可读存储介质，其上存储有计算机程序，其特征在于，所述计算机程序被处理器执行时实现权利要求1至8任一项所述的方法。"
        ),
    ]
    for claim in claims:
        add_para(doc, claim, size=12, first_line=False)


def add_description(doc: Document):
    add_heading(doc, "说明书")
    add_heading(doc, TITLE)

    add_section_title(doc, "技术领域")
    add_para(doc, "本发明涉及智能制造调度、共享制造资源优化和多目标组合优化技术领域，具体涉及一种面向多服务价值链协同的共享制造分布式柔性作业车间多目标调度方法及系统。")

    add_section_title(doc, "背景技术")
    add_para(doc, "共享制造通过第三方平台连接制造需求方和分散制造资源，使不同企业、车间或服务资源单元能够在统一机制下完成订单加工。与传统单车间或单企业调度相比，共享制造调度不仅需要考虑工序加工顺序、机器选择和完工时间，还需要处理订单制造需求与制造服务能力之间的匹配关系。")
    add_para(doc, "现有共享制造分布式柔性作业车间调度通常将制造服务资源单元抽象为SRU，并基于服务类型完成订单到SRU的供需匹配，再在SRU内部进行机器选择与工序排序。该类方法能够描述“服务类型-SRU-机器”的资源选择过程，但当第三方共享制造平台同时接入多条服务价值链时，仍难以区分订单的业务归属和订单的制造能力需求。")
    add_para(doc, "在实际制造平台中，价值链表示订单来源、客户归属、核心企业或产品族之间的协同关系；服务类型表示订单所需制造能力；SRU表示可被调度的制造服务资源单元；机器则是SRU内部具体执行工序加工的资源。若仅以服务类型作为调度依据，调度系统难以表达链内资源优先、跨链资源调用、跨链协同成本、跨链运输时间以及跨链资源缓解瓶颈负载等复杂关系。")
    add_para(doc, "同时，分布式柔性作业车间调度属于典型组合优化问题，在多目标场景中需要同时平衡总成本、最大完工时间和资源负载。通用多目标进化算法虽然具有较强适用性，但对价值链归属、服务类型匹配、链内/跨链资源选择和SRU内部车间调度之间的耦合关系利用不足，容易出现可行性修复频繁、局部搜索方向弱和非支配解分布不稳定等问题。")
    add_para(doc, "因此，亟需一种能够显式建模“价值链-服务类型-SRU-机器”层级关系，并能在链内资源和跨链资源之间进行多目标优化选择的共享制造调度方法。")

    add_section_title(doc, "发明内容")
    add_para(doc, "本发明所解决的技术问题在于提供一种面向多服务价值链协同的共享制造分布式柔性作业车间多目标调度方法及系统，用以在多条服务价值链共同接入共享制造平台的场景下，同时完成订单价值链归属识别、服务类型匹配、链内/跨链SRU选择、SRU内部机器选择和工序排序，并输出兼顾成本、时间和负载的非支配调度方案。")
    add_para(doc, "本发明提供的基础方案为：一种面向多服务价值链协同的共享制造分布式柔性作业车间多目标调度方法，包括以下步骤：")
    steps = [
        "S1、获取订单集合J、服务价值链集合V、服务类型集合T、SRU集合U、每个SRU内部机器集合Mu，以及订单工序集合Oj。对于每个订单，获取其价值链归属vcj、服务类型typej、释放时间、工序顺序和候选加工机器；对于每个SRU，获取其价值链归属vcu、可服务类型集合typesu、运输时间、运输成本和跨链固定协同成本。",
        "S2、构造候选SRU集合。对于订单j，先筛选满足typej属于typesu的SRU；若SRU与订单属于同一服务价值链，则加入链内候选集合；若SRU与订单属于不同服务价值链且跨链模式允许，则加入跨链候选集合。",
        "S3、生成四层编码。UA层表示订单到SRU的分配关系；OS层表示按服务类型组织的工序排序序列；OP层由UA和OS推导得到，表示每个SRU内部的工序加工序列；MS层表示每道工序在所选SRU内部的机器选择。",
        "S4、执行可行性修复。若UA层违反服务类型匹配或跨链模式，则重新选择可行SRU；若OS层工序令牌数量不满足订单工序数量，则修复OS层；若MS层机器不可加工对应工序，则从可行机器集合中重新选择。",
        "S5、解码评价。按照OP层中各SRU的工序序列和MS层的机器选择，维护订单就绪时间和机器就绪时间，计算每道工序开始时间与结束时间，并累计加工成本、运输成本、跨链协同成本和SRU负载。",
        "S6、形成多目标函数值，至少包括总成本F1和最大完工时间F2，优选地还包括最大SRU负载F3。",
        "S7、执行概率学习。将非支配解和精英解作为学习样本，更新订单分配概率矩阵、工序排序概率矩阵和机器选择概率矩阵。",
        "S8、执行局部搜索。采用链内SRU替换、跨链SRU替换、跨链回流、瓶颈SRU释放、工序插入和机器替换等邻域对候选解进行改进。",
        "S9、更新非支配解记忆池，重复迭代直至达到最大迭代次数或运行时间上限，输出非支配调度方案集合。"
    ]
    for step in steps:
        add_para(doc, step, first_line=False)

    add_para(doc, "在本发明中，候选SRU集合可表示为：")
    add_formula(doc, "Aj={u | typej ∈ typesu, shareu=1, and mode constraint is satisfied}")
    add_formula(doc, "Aj_in={u ∈ Aj | vcj = vcu}, Aj_cross={u ∈ Aj | vcj ≠ vcu}")
    add_para(doc, "其中，shareu表示SRU是否开放共享；在一种优选实施例中，所有SRU均开放共享，跨链是否可用由调度模式控制。")

    add_para(doc, "总成本F1优选由加工成本PC、运输成本TC和跨链固定协同成本CFC组成：")
    add_formula(doc, "F1 = PC + TC + CFC")
    add_formula(doc, "PC = ΣjΣoΣuΣm pj,o,u,m · cj,o,u,m · yj,o,u,m")
    add_formula(doc, "TC = ΣjΣu tcj,u · xj,u")
    add_formula(doc, "CFC = ΣjΣu fcj,u · qj,u")
    add_para(doc, "其中，xj,u表示订单j是否分配至SRU u，yj,o,u,m表示订单j的工序o是否在SRU u的机器m上加工，qj,u表示订单j是否跨服务价值链选择SRU u。当前优选实施例采用固定跨链协同成本；在其他实施方式中，也可根据管理需要增设跨链变动成本项。")
    add_formula(doc, "F2 = maxj(Cj + ttj,u)")
    add_formula(doc, "F3 = maxu(loadu), loadu = ΣjΣoΣm pj,o,u,m · yj,o,u,m")

    add_para(doc, "相比于现有技术，本发明至少具有以下有益效果：")
    benefits = [
        "1. 本发明显式区分价值链归属和服务类型需求，建立“价值链-服务类型-SRU-机器”的层级调度结构，使共享制造平台能够同时表达业务协同关系和制造能力匹配关系。",
        "2. 本发明将链内资源选择和跨链资源选择统一纳入候选SRU集合，并通过跨链固定协同成本、运输成本和跨链模式约束刻画跨服务价值链调度行为。",
        "3. 本发明采用UA、OS、OP、MS四层编码，使订单分配、工序排序、SRU内部工序序列和机器选择之间的映射关系清晰，可降低解码和可行性修复复杂度。",
        "4. 本发明通过概率模型学习优质解中的订单分配、工序排序和机器选择模式，并结合非支配解记忆池保持Pareto解的多样性。",
        "5. 本发明设计链内替换、跨链替换、跨链回流和瓶颈资源释放等问题特定邻域，使算法能够直接搜索成本、工期和负载之间的关键折中区域。",
        "6. 本发明可输出非支配调度方案集合，便于平台根据成本优先、工期优先或负载均衡优先等不同策略选择最终执行方案。"
    ]
    for benefit in benefits:
        add_para(doc, benefit, first_line=False)

    add_section_title(doc, "附图说明")
    figures = [
        "图1为本发明实施例的面向多服务价值链协同的共享制造分布式柔性作业车间多目标调度方法流程示意图；",
        "图2为本发明实施例的“价值链-服务类型-SRU-机器”层级调度结构示意图；",
        "图3为本发明实施例的四层编码及解码评价流程示意图；",
        "图4为本发明实施例的共享制造多目标调度系统结构示意图。"
    ]
    for fig in figures:
        add_para(doc, fig, first_line=False)

    add_section_title(doc, "具体实施方式")
    add_para(doc, "下面结合具体实施方式对本发明作进一步说明。应当理解，以下实施方式仅用于说明本发明，而不用于限定本发明的保护范围。")
    add_para(doc, "实施例基本如图1至图4所示。本实施例考虑一个第三方共享制造平台，该平台连接多条服务价值链、多个SRU及每个SRU内部的多台机器。以航空零部件共享制造为例，平台可连接发动机零部件制造链、机体结构件制造链和维修备件快速响应链。订单按照业务来源归属于不同服务价值链，并按照制造能力需求划分为不同服务类型。")
    add_para(doc, "平台接收到订单池后，首先读取订单的价值链归属和服务类型，再根据服务类型匹配关系确定候选SRU。若调度模式不允许跨链，则订单只能选择本价值链内且服务类型匹配的SRU；若调度模式允许跨链，则订单还可以选择其他价值链内且服务类型匹配的SRU。无论链内还是跨链，服务类型匹配均为硬约束。")
    add_para(doc, "在一个具体实施例中，扩展实例包括30个订单、3条服务价值链、2种服务类型、6个SRU、10类基础机器和231道工序。服务类型T1对应候选SRU U1、U3和U5，服务类型T2对应候选SRU U2、U4和U6。订单在3条服务价值链中均衡分布，各SRU均开放共享，但跨链选择会产生固定协同成本。")
    add_para(doc, "算法初始化时，先建立四层编码。UA层以键值形式记录订单编号和SRU编号的对应关系；OS层对同服务类型订单的工序令牌进行排序；OP层根据UA层和OS层将工序映射到各SRU内部序列；MS层记录每个SRU内部工序对应的机器编号。")
    add_para(doc, "在解码过程中，系统依次遍历各SRU内部的工序序列。对于每道工序，若MS层指定机器可加工该工序，则使用该机器；若指定机器不可行，则从可行机器中选择加工时间最短或加工成本最低的机器。系统将该工序的开始时间确定为订单就绪时间与机器就绪时间的较大值，并将结束时间确定为开始时间与加工时间之和。")
    add_para(doc, "解码完成后，系统计算调度方案的加工成本、运输成本、跨链固定协同成本、最大完工时间、各SRU负载、跨链订单数量、跨链比例、价值链流入流出关系和负载标准差等诊断信息。上述诊断信息可用于解释调度方案为何选择链内资源或跨链资源。")
    add_para(doc, "概率学习阶段，系统维护三类概率矩阵。PMA矩阵学习订单分配至不同候选SRU的概率；PMS矩阵学习订单工序令牌在排序序列中的位置概率；PMM矩阵学习每道工序在对应SRU内部选择不同机器的概率。每轮迭代中，系统基于精英解和非支配解记忆池更新概率矩阵，从而强化已发现的优质调度结构。")
    add_para(doc, "局部搜索阶段，系统生成多个邻域解。链内替换邻域用于在同一价值链的同服务类型SRU之间移动订单；跨链替换邻域用于将订单移动至其他价值链中同服务类型SRU；跨链回流邻域用于将已经跨链的订单迁回本价值链资源；瓶颈资源释放邻域用于从当前负载最高SRU迁出部分订单；工序插入邻域用于改变OS层中工序令牌位置；机器替换邻域用于改变MS层中工序加工机器。")
    add_para(doc, "局部搜索得到的邻域解通过可行性修复和解码评价后，与当前种群共同参与非支配排序。系统利用拥挤距离保持解集分布，并将非支配解写入非支配解记忆池。当记忆池超过容量时，系统保留分布更分散的调度方案。")
    add_para(doc, "最终，系统输出非支配调度方案集合。平台可根据实际策略从中选择折中方案，例如选择总成本最低方案、最大完工时间最短方案、最大SRU负载最低方案，或者选择经归一化加权距离计算得到的折中方案。")
    add_para(doc, "本发明还提供一种共享制造多目标调度系统。该系统包括数据获取模块、候选资源构造模块、编码初始化模块、可行性修复模块、解码评价模块、概率学习模块、局部搜索模块和输出模块。上述模块可以由软件程序实现，也可以由处理器、存储器和通信接口构成的硬件设备实现，或者由软件和硬件结合实现。")
    add_para(doc, "以上所述仅为本发明的较佳实施例，并不用于限制本发明。对于本领域普通技术人员而言，在不脱离本发明构思的前提下，可以对服务价值链数量、服务类型数量、SRU数量、目标函数组合、概率模型更新方式和局部搜索邻域进行若干变形和改进，这些变形和改进均应属于本发明的保护范围。")


def add_drawings(doc: Document):
    add_page_break(doc)
    add_heading(doc, "说明书附图")
    add_para(doc, "图1", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "订单池 -> 候选SRU构造 -> 四层编码初始化 -> 可行性修复 -> 解码评价 -> 概率模型更新 -> 局部搜索 -> 非支配解输出", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "图2", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "服务价值链层 -> 服务类型层 -> SRU资源层 -> SRU内部机器层", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "图3", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "UA订单-SRU分配；OS工序排序；OP按SRU生成工序序列；MS机器选择；解码后得到成本、完工时间和负载。", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "图4", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "数据获取模块、候选资源构造模块、编码初始化模块、可行性修复模块、解码评价模块、概率学习模块、局部搜索模块和输出模块。", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, OUT_DOCX)
    doc = Document(OUT_DOCX)
    clear_body(doc)
    set_doc_defaults(doc)

    add_heading(doc, "说明书摘要")
    abstract = (
        "本发明涉及智能制造调度技术领域，具体涉及一种面向多服务价值链协同的共享制造分布式柔性作业车间多目标调度方法及系统。"
        "该方法获取订单集合、服务价值链集合、服务类型集合、制造服务资源单元集合及其内部机器集合，依据订单价值链归属和服务类型构造链内与跨链候选资源集合；"
        "建立UA、OS、OP和MS四层编码，对订单-SRU分配、工序排序、SRU内部工序序列和机器选择进行统一表达；"
        "通过可行性修复保证服务类型匹配、跨链模式、工艺顺序和机器加工约束；"
        "解码计算加工成本、运输成本、跨链固定协同成本、最大完工时间和SRU负载；"
        "再利用非支配解记忆池、概率模型更新以及链内替换、跨链替换、跨链回流、瓶颈资源释放等局部搜索获得非支配调度方案集合。"
        "本发明能够在多服务价值链共享制造场景下同时刻画业务归属、制造能力匹配和车间调度过程，提高链内/跨链资源协同决策的可计算性和可解释性。"
    )
    add_para(doc, abstract)
    add_para(doc, "")
    add_heading(doc, "摘要附图")
    add_para(doc, "图1", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "订单池 -> 候选SRU构造 -> 四层编码 -> 解码评价 -> 概率学习与局部搜索 -> 非支配调度方案", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_page_break(doc)

    add_claims(doc)
    add_page_break(doc)
    add_description(doc)
    add_drawings(doc)

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    path = build()
    print(path)
