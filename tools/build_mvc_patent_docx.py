from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "sm_dfjsp" / "docs"
ASSET_DIR = OUT_DIR / "patent_assets_mvc_smdfjsp"
OUT = OUT_DIR / "MVC_SM_DFJSP_patent_draft.docx"


def _build_figures() -> tuple[Path, Path]:
    from PIL import Image, ImageDraw, ImageFont

    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    font_candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    font_path = next((p for p in font_candidates if p.exists()), None)
    font = ImageFont.truetype(str(font_path), 28) if font_path else ImageFont.load_default()
    small = ImageFont.truetype(str(font_path), 22) if font_path else ImageFont.load_default()
    title = ImageFont.truetype(str(font_path), 32) if font_path else ImageFont.load_default()

    def centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt) -> None:
        lines = text.split("\n")
        sizes = []
        for line in lines:
            b = draw.textbbox((0, 0), line, font=fnt)
            sizes.append((b[2] - b[0], b[3] - b[1]))
        total_h = sum(h for _, h in sizes) + (len(lines) - 1) * 8
        y = box[1] + (box[3] - box[1] - total_h) / 2
        for line, (w, h) in zip(lines, sizes):
            x = box[0] + (box[2] - box[0] - w) / 2
            draw.text((x, y), line, font=fnt, fill=(0, 0, 0))
            y += h + 8

    def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
        import math

        draw.line([start, end], fill=(45, 45, 45), width=3)
        x1, y1 = start
        x2, y2 = end
        ang = math.atan2(y2 - y1, x2 - x1)
        for a in (ang + 2.55, ang - 2.55):
            draw.line([end, (x2 + 14 * math.cos(a), y2 + 14 * math.sin(a))], fill=(45, 45, 45), width=3)

    img = Image.new("RGB", (1400, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((420, 40), "图1  调度方法流程示意图", font=title, fill=(0, 0, 0))
    boxes = [
        (80, 130, 420, 230, "S1 获取订单、价值链、\n服务类型和资源数据"),
        (530, 130, 870, 230, "S2 构建链内/跨链\n候选 SRU 集合"),
        (980, 130, 1320, 230, "S3 建立双目标\n调度评价模型"),
        (980, 350, 1320, 450, "S4 生成 UA/OS/OP/MS\n四层编码种群"),
        (530, 350, 870, 450, "S5 价值链先验 PMA\n与精英频率融合"),
        (80, 350, 420, 450, "S6 采样、修复\n并解码调度"),
        (80, 570, 420, 670, "S7 跨链协同\n禁忌搜索"),
        (530, 570, 870, 670, "S8 更新非支配档案\n与邻域概率"),
        (980, 570, 1320, 670, "S9 输出 Pareto 解集、\n调度方案和诊断指标"),
    ]
    for x1, y1, x2, y2, text in boxes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=8, outline=(20, 73, 112), width=3, fill=(235, 243, 248))
        centered(d, (x1, y1, x2, y2), text, font)
    for start, end in [
        ((420, 180), (530, 180)),
        ((870, 180), (980, 180)),
        ((1150, 230), (1150, 350)),
        ((980, 400), (870, 400)),
        ((530, 400), (420, 400)),
        ((250, 450), (250, 570)),
        ((420, 620), (530, 620)),
        ((870, 620), (980, 620)),
    ]:
        arrow(d, start, end)
    fig1 = ASSET_DIR / "fig1_method_flow.png"
    img.save(fig1)

    img2 = Image.new("RGB", (1400, 820), "white")
    d = ImageDraw.Draw(img2)
    d.text((390, 35), "图2  系统模块结构示意图", font=title, fill=(0, 0, 0))
    modules = [
        (80, 150, 330, 270, "数据获取模块"),
        (410, 150, 660, 270, "候选资源构建模块"),
        (740, 150, 990, 270, "模型评价模块"),
        (1070, 150, 1320, 270, "结果输出模块"),
        (245, 420, 495, 540, "概率学习模块"),
        (575, 420, 825, 540, "局部搜索模块"),
        (905, 420, 1155, 540, "档案管理模块"),
    ]
    for x1, y1, x2, y2, text in modules:
        d.rounded_rectangle((x1, y1, x2, y2), radius=8, outline=(0, 0, 0), width=3, fill=(248, 248, 248))
        centered(d, (x1, y1, x2, y2), text, font)
    for start, end in [
        ((330, 210), (410, 210)),
        ((660, 210), (740, 210)),
        ((990, 210), (1070, 210)),
        ((865, 270), (705, 420)),
        ((495, 480), (575, 480)),
        ((825, 480), (905, 480)),
        ((1030, 420), (1195, 270)),
        ((705, 540), (705, 620)),
        ((705, 620), (410, 620)),
        ((410, 620), (410, 540)),
    ]:
        arrow(d, start, end)
    d.rounded_rectangle((160, 650, 1240, 750), radius=8, outline=(20, 73, 112), width=3, fill=(235, 243, 248))
    centered(d, (160, 650, 1240, 750), "共享制造平台数据库：订单、价值链、服务类型、SRU、机器、加工时间、成本和运输参数", small)
    fig2 = ASSET_DIR / "fig2_system_modules.png"
    img2.save(fig2)
    return fig1, fig2


def _configure_styles(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.05)
        section.right_margin = Inches(1.05)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.35
    styles["Normal"].paragraph_format.first_line_indent = Pt(24)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = styles[name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)

    st = styles.add_style("Patent Title", WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = "黑体"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    st.font.size = Pt(18)
    st.font.bold = True
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(12)

    st = styles.add_style("Claim", WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = "宋体"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    st.font.size = Pt(12)
    st.paragraph_format.line_spacing = 1.35
    st.paragraph_format.first_line_indent = Pt(0)
    st.paragraph_format.space_after = Pt(6)


def _add_para(doc: Document, text: str, style: str | None = None, first_indent: bool = True):
    p = doc.add_paragraph(style=style)
    if not first_indent:
        p.paragraph_format.first_line_indent = Pt(0)
    p.add_run(text)
    return p


def _add_fig(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.5))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.first_line_indent = Pt(0)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)
            if r == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "D9EAF7")
                cell._tc.get_or_add_tcPr().append(shading)


def build_docx() -> Path:
    fig1, fig2 = _build_figures()
    doc = Document()
    _configure_styles(doc)

    title = "一种面向多服务价值链协同的共享制造分布式柔性作业车间调度方法及系统"

    doc.add_paragraph("说明书摘要", style="Heading 1")
    _add_para(
        doc,
        "本发明涉及共享制造与智能调度技术领域，具体涉及一种面向多服务价值链协同的共享制造分布式柔性作业车间调度方法及系统。该方法包括：S1、获取共享制造平台中的订单、价值链、服务类型、服务资源单元、机器、加工时间、加工成本、运输时间、运输成本以及跨链协调成本数据；S2、根据订单所属价值链和服务类型构建链内候选服务资源单元集合与跨链候选服务资源单元集合；S3、在允许跨链和不允许跨链两种模式下建立以总成本和最大完工时间最小化为目标的调度评价模型；S4、采用订单资源分配、工序排序、服务资源单元内部工序队列和机器选择四层编码生成候选调度方案；S5、基于加工成本、运输成本、跨链固定协调成本、预计完成时间和跨链时间收益构建价值链感知先验概率，并与精英解频率融合更新概率模型；S6、对采样个体进行可行性修复和解码；S7、采用链内替换、跨链替换、跨链回流、关键订单迁移、高成本跨链回流以及机器与工序重排的禁忌搜索进行局部优化；S8、更新非支配解档案并输出 Pareto 调度方案及跨链比例、价值链流入流出和服务资源单元负载等诊断指标。本发明能够在供需类型匹配的基础上显式区分链内协同与跨链协同，提高多价值链共享制造场景下资源调用、成本控制和工期压缩的综合决策能力。",
    )
    doc.add_paragraph("摘要附图", style="Heading 1")
    _add_fig(doc, fig1, "摘要附图")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("权利要求书", style="Heading 1")
    claims = [
        "1. 一种面向多服务价值链协同的共享制造分布式柔性作业车间调度方法，其特征在于，包括以下步骤：S1、获取共享制造平台的订单集合、价值链集合、服务类型集合、服务资源单元集合、机器集合、加工时间、加工成本、运输时间、运输成本以及跨链协调成本数据；S2、依据订单所属价值链和订单服务类型，为每个订单构建链内候选服务资源单元集合和跨链候选服务资源单元集合；S3、根据跨链模式确定订单可选资源范围，并建立以总成本最小化和最大完工时间最小化为目标的调度评价模型；S4、生成包含订单资源分配 UA、工序排序 OS、服务资源单元内部工序队列 OP 和机器选择 MS 的四层编码个体；S5、基于价值链感知概率模型采样生成新个体，并对不满足服务类型匹配或跨链模式约束的个体进行可行性修复；S6、对修复后的个体进行解码，计算加工成本、运输成本、跨链固定协调成本、最大完工时间以及跨链诊断指标；S7、对解码后的个体执行跨链协同禁忌搜索，获得局部改进解；S8、基于 Pareto 支配关系更新非支配解档案，并输出调度方案。",
        "2. 根据权利要求1所述的调度方法，其特征在于，所述链内候选服务资源单元集合由服务类型与订单需求匹配且所属价值链与订单所属价值链相同的服务资源单元构成，所述跨链候选服务资源单元集合由服务类型与订单需求匹配且所属价值链与订单所属价值链不同的服务资源单元构成。",
        "3. 根据权利要求1所述的调度方法，其特征在于，当跨链模式为不允许跨链时，订单仅能从链内候选服务资源单元集合中选择服务资源单元；当跨链模式为允许跨链时，订单能从链内候选服务资源单元集合与跨链候选服务资源单元集合的并集中选择服务资源单元。",
        "4. 根据权利要求1所述的调度方法，其特征在于，所述总成本包括所有订单工序的加工成本、订单由服务资源单元送达需求节点的运输成本以及订单跨价值链调用服务资源单元时产生的跨链固定协调成本。",
        "5. 根据权利要求1所述的调度方法，其特征在于，所述最大完工时间由订单最后一道工序完成时间与该订单所选服务资源单元到需求节点的运输时间共同确定。",
        "6. 根据权利要求1所述的调度方法，其特征在于，所述四层编码中，UA 层用于表示每个订单选择的服务资源单元，OS 层用于表示工序加工顺序，OP 层由 UA 层和 OS 层推导得到并用于表示每个服务资源单元内部的工序队列，MS 层用于表示服务资源单元内部每道工序选择的机器。",
        "7. 根据权利要求1所述的调度方法，其特征在于，所述价值链感知概率模型包括订单资源分配概率矩阵、工序排序概率矩阵和机器选择概率矩阵，其中订单资源分配概率矩阵由精英解中的订单-服务资源单元选择频率与价值链感知先验概率融合更新。",
        "8. 根据权利要求7所述的调度方法，其特征在于，所述价值链感知先验概率根据候选服务资源单元对应的加工成本、运输成本、跨链固定协调成本、预计完成时间以及相对于链内候选服务资源单元的跨链时间收益计算得到。",
        "9. 根据权利要求1所述的调度方法，其特征在于，所述跨链协同禁忌搜索至少包括链内服务资源单元替换、跨链服务资源单元替换、跨链回流、关键订单迁移、高成本跨链回流以及机器与工序重排中的一种或多种邻域操作。",
        "10. 根据权利要求9所述的调度方法，其特征在于，所述跨链协同禁忌搜索记录各邻域操作的接受次数、进入非支配解档案次数和目标改进幅度，并根据记录结果自适应更新各邻域操作的选择概率。",
        "11. 根据权利要求1所述的调度方法，其特征在于，所述非支配解档案用于保存搜索过程中获得的 Pareto 非支配解，并作为概率模型更新的学习样本和禁忌搜索的局部搜索种子。",
        "12. 根据权利要求1所述的调度方法，其特征在于，所述跨链诊断指标包括跨链订单比例、价值链间跨链调用流、价值链流入订单数、价值链流出订单数、服务资源单元负载标准差以及最大服务资源单元负载中的一种或多种。",
        "13. 一种应用权利要求1至12任一项所述调度方法的共享制造分布式柔性作业车间调度系统，其特征在于，包括：数据获取模块，用于获取订单、价值链、服务类型、服务资源单元、机器、加工时间、加工成本、运输时间、运输成本和跨链协调成本数据；候选资源构建模块，用于构建链内候选服务资源单元集合和跨链候选服务资源单元集合；模型评价模块，用于根据总成本和最大完工时间评价调度方案；概率学习模块，用于更新订单资源分配概率矩阵、工序排序概率矩阵和机器选择概率矩阵；局部搜索模块，用于执行跨链协同禁忌搜索；档案管理模块，用于维护非支配解档案；结果输出模块，用于输出 Pareto 调度方案和跨链诊断指标。",
        "14. 一种电子设备，其特征在于，包括处理器和存储器，所述存储器中存储有计算机程序，所述计算机程序被所述处理器执行时实现权利要求1至12任一项所述的调度方法。",
        "15. 一种计算机可读存储介质，其特征在于，所述计算机可读存储介质中存储有计算机程序，所述计算机程序被处理器执行时实现权利要求1至12任一项所述的调度方法。",
    ]
    for claim in claims:
        _add_para(doc, claim, style="Claim", first_indent=False)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("说明书", style="Heading 1")
    doc.add_paragraph(title, style="Patent Title")

    sections = [
        (
            "技术领域",
            [
                "本发明涉及共享制造、分布式柔性作业车间调度、多目标优化和智能调度算法技术领域，具体涉及一种面向多服务价值链协同的共享制造分布式柔性作业车间调度方法及系统。"
            ],
        ),
        (
            "背景技术",
            [
                "共享制造平台通过网络化方式聚合分散制造资源，使订单能够在不同服务资源单元中获得加工服务。分布式柔性作业车间调度问题需要同时决定订单分配、工序排序和机器选择，既要满足工序先后关系、机器能力约束和资源可用性约束，又要兼顾成本、交期和资源利用效率。",
                "现有共享制造分布式柔性作业车间调度研究通常以供需类型匹配为基础，即订单具有某类制造服务需求，服务资源单元提供相应服务类型，调度系统在类型匹配的服务资源单元中选择加工资源。该类方法能够解决制造服务需求与资源能力的一致性问题，但通常没有显式刻画多个服务价值链共存时的资源边界和协同方式。",
                "在第三方共享制造平台中，不同核心企业、不同客户群或不同业务网络可形成多条服务价值链。每条价值链内部拥有一定服务资源，同时部分资源可以向其他价值链开放共享。若仅按服务类型匹配进行调度，系统难以区分订单使用本价值链资源还是跨价值链资源，也难以量化跨链调用带来的协调成本、运输成本、交付时间变化和负载转移效果。",
                "此外，传统估计分布算法与禁忌搜索结合的 EDA-TS 方法主要面向一般资源分配、工序排序和机器选择，其概率学习通常根据精英解频率更新，缺少对价值链归属、链内候选资源和跨链候选资源的结构化建模。当跨链资源具有时间优势或链内资源拥堵时，现有算法可能无法有效发现跨链协同带来的工期改善；当跨链调用只增加协调成本时，现有算法也缺少将订单回流到链内资源的针对性搜索机制。",
                "因此，有必要提供一种能够在共享制造分布式柔性作业车间调度中显式表达多服务价值链、服务类型匹配、链内协同、跨链协同和多目标优化的调度方法及系统。",
            ],
        ),
        (
            "发明内容",
            [
                "本发明所要解决的技术问题在于提供一种面向多服务价值链协同的共享制造分布式柔性作业车间调度方法及系统，以解决现有方法不能同时表达价值链归属、供需类型匹配和跨链资源调用，以及传统 EDA-TS 算法缺少价值链感知概率学习和跨链协同邻域的问题。",
                "一种面向多服务价值链协同的共享制造分布式柔性作业车间调度方法，包括以下步骤：S1、数据获取；S2、候选资源构建；S3、跨链模式设置；S4、调度模型建立；S5、个体编码；S6、概率学习；S7、采样与修复；S8、解码评价；S9、跨链协同禁忌搜索；S10、档案更新与结果输出。",
                "S1 中，获取共享制造平台的订单集合、价值链集合、服务类型集合、服务资源单元集合、机器集合、订单工序集合、工序加工时间、单位加工成本、订单至服务资源单元的运输时间和运输成本、服务资源单元所属价值链、服务资源单元可提供服务类型以及跨链固定协调成本。",
                "S2 中，对每个订单，根据订单所属价值链和服务类型构建候选资源集合。属于同一价值链且服务类型匹配的服务资源单元构成链内候选集合；属于其他价值链且服务类型匹配的服务资源单元构成跨链候选集合。",
                "S3 中，设置不允许跨链模式和允许跨链模式。不允许跨链模式下，订单只能选择链内候选服务资源单元；允许跨链模式下，订单能够选择链内候选服务资源单元或跨链候选服务资源单元。",
                "S4 中，建立双目标评价模型。第一目标为总成本最小化，总成本包括加工成本、运输成本和跨链固定协调成本；第二目标为最大完工时间最小化，最大完工时间由订单最后一道工序完成时间和运输时间共同确定。",
                "S5 中，构造 UA、OS、OP 和 MS 四层编码，其中 UA 表示订单到服务资源单元的分配，OS 表示工序排序，OP 表示每个服务资源单元内部的工序队列，MS 表示每道工序在已选服务资源单元内的机器选择。",
                "S6 中，构建订单资源分配概率矩阵、工序排序概率矩阵和机器选择概率矩阵。订单资源分配概率矩阵融合精英解频率和价值链感知先验概率；价值链感知先验概率由候选服务资源单元对应的成本、时间、跨链固定协调成本和跨链时间收益确定。",
                "S7 中，依据概率模型采样生成新个体。若新个体违反服务类型匹配约束、跨链模式约束或机器可加工约束，则在对应候选集合中进行修复，并重新生成服务资源单元内部工序队列。",
                "S8 中，对可行个体进行解码，维护订单就绪时间和机器就绪时间，计算每道工序的开始时间和完成时间，并得到总成本、最大完工时间、服务资源单元负载、跨链比例和价值链流向。",
                "S9 中，对个体执行链内服务资源单元替换、跨链服务资源单元替换、跨链回流、关键订单迁移、高成本跨链回流以及机器与工序重排等邻域搜索，并利用禁忌表避免搜索循环。",
                "S10 中，根据 Pareto 支配关系更新非支配解档案，将非支配解档案作为后续概率学习样本和局部搜索种子，最终输出 Pareto 调度方案、成本-工期折中结果和跨链诊断指标。",
                "相比于现有技术，本发明能够在共享制造调度中引入价值链归属层，使订单选择资源时不仅考虑服务类型匹配，还考虑链内资源和跨链资源的差异；能够通过 cross-off 与 cross-on 两种模式比较不允许跨链和允许跨链时的成本、工期和资源流动变化；能够通过价值链感知概率模型和跨链协同邻域提高多价值链共享制造场景中的资源调用和调度优化能力。",
            ],
        ),
        (
            "附图说明",
            [
                "图1为本发明实施例中一种面向多服务价值链协同的共享制造分布式柔性作业车间调度方法的流程示意图。",
                "图2为本发明实施例中一种面向多服务价值链协同的共享制造分布式柔性作业车间调度系统的模块结构示意图。",
            ],
        ),
        (
            "具体实施方式",
            [
                "下面结合附图和实施例对本发明作进一步说明。以下实施例用于说明本发明的技术方案，并非用于限定本发明的保护范围。",
                "实施例一：如图1所示，共享制造平台接收多个订单。每个订单具有所属价值链和服务类型，每个服务资源单元具有所属价值链、可提供服务类型以及内部机器集合。平台首先读取订单、服务资源单元、机器、加工时间、加工成本、运输时间、运输成本和跨链固定协调成本等数据。",
                "对于订单 j，平台根据服务类型匹配规则筛选候选服务资源单元。若候选服务资源单元 u 的可提供服务类型包含订单 j 的服务类型，且 u 的所属价值链与订单 j 的所属价值链相同，则 u 被加入链内候选集合；若候选服务资源单元 u 的可提供服务类型包含订单 j 的服务类型，且 u 的所属价值链与订单 j 的所属价值链不同，则 u 被加入跨链候选集合。",
                "在不允许跨链模式下，订单 j 的可选服务资源单元集合为链内候选集合；在允许跨链模式下，订单 j 的可选服务资源单元集合为链内候选集合与跨链候选集合的并集。通过该方式，系统能够使用同一数据结构比较链内调度方案和跨链协同调度方案。",
                "调度个体采用四层编码。UA 层记录每个订单分配到哪个服务资源单元，是表达链内协同和跨链协同的关键层；OS 层记录订单工序排序；OP 层由 UA 层和 OS 层推导得到，用于描述各服务资源单元内部的工序队列；MS 层记录每道工序在所选服务资源单元内部选择的机器。",
                "在概率学习阶段，系统分别维护订单资源分配概率矩阵、工序排序概率矩阵和机器选择概率矩阵。订单资源分配概率矩阵的先验评分可由加工成本、运输成本、跨链固定协调成本、预计完成时间和跨链时间收益加权得到，评分越低表示候选服务资源单元越优。系统将该先验概率与精英解中订单选择服务资源单元的频率融合，形成新的目标概率，再按学习率更新订单资源分配概率矩阵。",
                "在局部搜索阶段，系统执行跨链协同禁忌搜索。链内服务资源单元替换用于在同一价值链、同一服务类型的服务资源单元之间移动订单；跨链服务资源单元替换用于将订单迁移到其他价值链的同服务类型服务资源单元；跨链回流用于将已跨链订单迁回本价值链资源；关键订单迁移用于将影响最大完工时间的订单迁移到预计完成时间更短的候选资源；高成本跨链回流用于将跨链成本较高的订单迁回链内或低成本候选资源；机器与工序重排用于改善服务资源单元内部的柔性作业车间排程。",
                "系统使用非支配解档案保存 Pareto 优质调度方案。非支配解档案中的解既用于最终输出，也用于下一轮概率学习和局部搜索种子选择。当档案规模超过设定容量时，系统优先保留目标空间中分布较分散的解，以维持成本与工期折中的多样性。",
            ],
        ),
    ]
    for heading, paragraphs in sections:
        doc.add_paragraph(heading, style="Heading 2")
        for paragraph in paragraphs:
            _add_para(doc, paragraph)
        if heading == "具体实施方式":
            _add_fig(doc, fig2, "图2  系统模块结构示意图")

    doc.add_paragraph("实施例二：实验数据与参数", style="Heading 2")
    _add_para(
        doc,
        "在一个可选实施例中，采用由 MK/FJSP benchmark 扩展得到的 MVC-MK01 至 MVC-MK15 实例作为测试数据。该数据集设置两条价值链、两类服务类型和四个服务资源单元。每个订单具有一个服务类型匹配的链内服务资源单元和一个服务类型匹配的跨链服务资源单元，候选服务资源单元上的加工时间保持一致，跨链差异主要体现为运输参数和跨链固定协调成本。",
    )
    _add_para(
        doc,
        "在该实施例中，正式总成本为加工成本、运输成本和跨链固定协调成本之和；跨链变动成本率字段作为兼容字段保留但不计入目标函数。算法可设置为双目标优化，种群规模为 80，最大迭代次数为 150，随机种子为 20260428 至 20260432，时间上限为 12000 秒。对比算法可包括 NSGA-II、MOEA/D、普通 EDA-TS 和本发明的 MVC-EDA-TS，并分别在不允许跨链模式和允许跨链模式下运行。",
    )
    _add_table(
        doc,
        [
            ["项目", "示例设置"],
            ["实例", "MVC-MK01 至 MVC-MK15"],
            ["价值链数量", "2"],
            ["服务类型数量", "2"],
            ["服务资源单元数量", "4"],
            ["目标", "总成本最小化、最大完工时间最小化"],
            ["跨链模式", "不允许跨链、允许跨链"],
            ["种群规模", "80"],
            ["最大迭代次数", "150"],
            ["随机种子", "20260428 至 20260432"],
        ],
    )
    _add_para(doc, "上述实施例表明，本发明能够在同一共享制造调度框架中比较链内调度和跨链协同调度，分析跨链协同对成本、工期、价值链流向和服务资源单元负载的影响。")
    _add_para(doc, "以上所述仅为本发明的实施例。对于本领域普通技术人员而言，在不脱离本发明构思的前提下，还可以对价值链数量、服务类型数量、服务资源单元数量、跨链成本函数、邻域操作种类和概率学习参数作出若干变形或替换，这些变形或替换均应落入本发明的保护范围。")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_docx())
