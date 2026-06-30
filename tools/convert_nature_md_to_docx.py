from __future__ import annotations

import re
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
TABLE_FILL = "F4F6F9"
BORDER = "B7C3D0"
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, name: str = "Calibri", east_asia: str = "SimSun", size: float | None = None,
                 color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.208

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("MVC-SM-DFJSP manuscript draft")
    set_run_font(footer_run, size=9, color=MUTED)


def clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    return text.rstrip()


def latex_to_plain(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\\text\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\mathrm\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\hat\{([^{}]*)\}", r"hat(\1)", text)
    replacements = {
        r"\times": "×",
        r"\in": "∈",
        r"\notin": "∉",
        r"\ne": "≠",
        r"\neq": "≠",
        r"\ge": "≥",
        r"\le": "≤",
        r"\cup": "∪",
        r"\sum": "∑",
        r"\max": "max",
        r"\min": "min",
        r"\forall": "∀",
        r"\mid": "|",
        r"\quad": "  ",
        r"\lambda": "λ",
        r"\alpha": "α",
        r"\tau": "τ",
        r"\exp": "exp",
        r"\bmod": "mod",
        r"\_": "_",
        r"\{": "{",
        r"\}": "}",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    text = re.sub(r"\\([A-Za-z]+)", r"\1", text)
    return re.sub(r"\s+", " ", text).strip()


def add_math_run(paragraph, formula: str, base_size: float = 11) -> None:
    run = paragraph.add_run(latex_to_plain(formula))
    set_run_font(run, name="Cambria Math", east_asia="SimSun", size=base_size)


def add_inline_runs(paragraph, text: str, base_size: float = 11) -> None:
    text = clean_text(text)
    if not text:
        return
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\\\(.+?\\\)|(?<!\$)\$[^$\n]+\$(?!\$))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", east_asia="SimSun", size=base_size - 0.5)
        elif token.startswith(r"\("):
            add_math_run(paragraph, token[2:-2], base_size=base_size)
        elif token.startswith("$"):
            add_math_run(paragraph, token[1:-1], base_size=base_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)


def split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in cells)


def column_widths(rows: list[list[str]]) -> list[int]:
    if not rows:
        return [CONTENT_WIDTH_DXA]
    cols = max(len(row) for row in rows)
    weights = []
    for idx in range(cols):
        max_len = max((len(row[idx]) if idx < len(row) else 0) for row in rows)
        weights.append(max(8, min(max_len, 42)))
    total = sum(weights) or cols
    widths = [max(900, int(CONTENT_WIDTH_DXA * w / total)) for w in weights]
    diff = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += diff
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) >= 2 and is_table_separator("|" + "|".join(rows[1]) + "|"):
        rows = [rows[0]] + rows[2:]
    cols = max(len(row) for row in rows)
    widths = column_widths(rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_borders(table)
    set_table_width(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            text = row[c_idx] if c_idx < len(row) else ""
            add_inline_runs(p, text, base_size=9.2 if cols >= 5 else 10)
            if r_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
                for run in p.runs:
                    run.bold = True
            if re.fullmatch(r"[-+]?[\d,.]+(%|e[+-]?\d+)?", text.strip(), re.I):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("\n".join(lines))
    set_run_font(run, name="Consolas", east_asia="SimSun", size=9)


def add_formula_block(doc: Document, formula: str, formula_dir: Path, formula_index: int) -> None:
    formula = formula.strip()
    if not formula:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    max_line_len = max((len(line) for line in formula.splitlines()), default=len(formula))
    width_inches = min(6.2, max(2.2, max_line_len * 0.075))
    image_path = formula_dir / f"formula_{formula_index:03d}.png"
    try:
        from PIL import Image, ImageDraw, ImageFont

        plain = latex_to_plain(formula)
        wrapped_lines: list[str] = []
        for source_line in plain.splitlines() or [plain]:
            line = source_line.strip()
            while len(line) > 92:
                cut = max(line.rfind(" ", 0, 92), line.rfind("+", 0, 92), line.rfind("∑", 1, 92))
                if cut <= 12:
                    cut = 92
                wrapped_lines.append(line[:cut].strip())
                line = line[cut:].strip()
            wrapped_lines.append(line)
        font_paths = [
            Path(r"C:\Windows\Fonts\cambria.ttc"),
            Path(r"C:\Windows\Fonts\consola.ttf"),
            Path(r"C:\Windows\Fonts\simsun.ttc"),
        ]
        font_path = next((fp for fp in font_paths if fp.exists()), None)
        font = ImageFont.truetype(str(font_path), 30) if font_path else ImageFont.load_default()
        dummy = Image.new("RGBA", (1, 1), (255, 255, 255, 0))
        draw = ImageDraw.Draw(dummy)
        boxes = [draw.textbbox((0, 0), line, font=font) for line in wrapped_lines]
        text_w = max((box[2] - box[0] for box in boxes), default=600)
        line_h = max((box[3] - box[1] for box in boxes), default=34) + 12
        img_w = max(800, text_w + 80)
        img_h = max(70, line_h * len(wrapped_lines) + 34)
        img = Image.new("RGBA", (img_w, img_h), (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        y = 17
        for line, box in zip(wrapped_lines, boxes):
            line_w = box[2] - box[0]
            draw.text(((img_w - line_w) / 2, y), line, fill=(0, 0, 0, 255), font=font)
            y += line_h
        img.save(image_path)
        p.add_run().add_picture(str(image_path), width=Inches(width_inches))
        return
    except Exception:
        pass

    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        render_text = formula if "\n" not in formula else latex_to_plain(formula)
        use_math = "\n" not in formula
        fig_height = 0.42 + 0.22 * max(1, len(formula.splitlines()))
        fig_width = width_inches
        fig = plt.figure(figsize=(fig_width, fig_height), dpi=240)
        if use_math:
            fig.text(0.5, 0.5, f"${render_text}$", ha="center", va="center", fontsize=12)
        else:
            fig.text(0.5, 0.5, render_text, ha="center", va="center", fontsize=10.5, family="DejaVu Sans Mono")
        fig.savefig(image_path, transparent=True, bbox_inches="tight", pad_inches=0.08)
        plt.close(fig)
        p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    except Exception:
        run = p.add_run(latex_to_plain(formula))
        set_run_font(run, name="Cambria Math", east_asia="SimSun", size=10.5)


def add_title_block(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, name="Calibri", east_asia="SimHei", size=22, color=RGBColor(0, 0, 0), bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(26)
    r = meta.add_run("Nature-style manuscript draft | Converted from Markdown")
    set_run_font(r, size=10.5, color=MUTED, italic=True)


def convert(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_styles(doc)
    tmp_formula_dir = tempfile.TemporaryDirectory()
    formula_dir = Path(tmp_formula_dir.name)
    formula_index = 0

    title = None
    idx = 0
    while idx < len(lines):
        if lines[idx].startswith("# "):
            title = lines[idx][2:].strip()
            break
        idx += 1
    add_title_block(doc, title or md_path.stem)

    i = 0
    in_code = False
    code_lines: list[str] = []
    paragraph_buf: list[str] = []

    def flush_para() -> None:
        nonlocal paragraph_buf
        if paragraph_buf:
            para_text = " ".join(x.strip() for x in paragraph_buf if x.strip())
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline_runs(p, para_text)
            paragraph_buf = []

    def add_display_formula(formula: str) -> None:
        nonlocal formula_index
        formula_index += 1
        add_formula_block(doc, formula, formula_dir, formula_index)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_para()
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            flush_para()
            i += 1
            continue

        if stripped.startswith(r"\[") and stripped.endswith(r"\]") and len(stripped) > 4:
            flush_para()
            add_display_formula(stripped[2:-2].strip())
            i += 1
            continue

        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            flush_para()
            add_display_formula(stripped[2:-2].strip())
            i += 1
            continue

        if stripped in {r"\[", "$$"}:
            flush_para()
            end_token = r"\]" if stripped == r"\[" else "$$"
            formula_lines: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != end_token:
                formula_lines.append(lines[i])
                i += 1
            add_display_formula("\n".join(formula_lines))
            if i < len(lines):
                i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_para()
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            flush_para()
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1 and text == title:
                i += 1
                continue
            style = "Heading 1" if level <= 2 else "Heading 2" if level == 3 else "Heading 3"
            p = doc.add_paragraph(style=style)
            add_inline_runs(p, text, base_size=doc.styles[style].font.size.pt if doc.styles[style].font.size else 12)
            i += 1
            continue

        if stripped.startswith(">"):
            flush_para()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped.lstrip("> ").strip())
            set_run_font(run, size=10, color=MUTED, italic=True)
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_para()
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        paragraph_buf.append(stripped)
        i += 1

    flush_para()
    if code_lines:
        add_code_block(doc, code_lines)

    doc.save(out_path)
    tmp_formula_dir.cleanup()


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: convert_nature_md_to_docx.py input.md output.docx")
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
